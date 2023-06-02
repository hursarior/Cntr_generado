import docx


# Abrir el documento original
doc = docx.Document("contrato1.docx")

# Copiar el estilo del primer párrafo del documento original al primer párrafo del nuevo documento
doc.paragraphs[8].add_run("6950 COOLIDGE ST HOLLYWOOD, FL 33024-3818 ")
doc.paragraphs[9].add_run("514111112350")
doc.paragraphs[48].add_run("\n LLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLLL")
ver = doc.paragraphs[10]
print(ver)
doc.paragraphs.insert(8, "lago")
doc.add_paragraph(doc.paragraphs[47].text, style=doc.paragraphs[47].style)
doc.add_paragraph(doc.paragraphs[8].text, style=doc.paragraphs[8].style)
doc.add_paragraph(doc.paragraphs[9].text, style=doc.paragraphs[9].style)


# Guardar el nuevo documento
doc.save("copiado.docx")

