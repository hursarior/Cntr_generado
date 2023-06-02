import tkinter as tk
from tkinter import messagebox
import docx

doc = docx.Document("contrato1.docx")



    
ventana = tk.Tk()
ventana.title("Formulario De Contratos")
ventana.geometry("500x450")

ventana.resizable(False, False)



def limitar_caracteres(valor_ingresado):
    if len(valor_ingresado) > 8:
        return False
    return True

validacion = ventana.register(limitar_caracteres)


nombre_label = tk.Label(ventana, text="DATOS DEL CLIENTE",font="helvetica 16")
nombre_label.grid(row=0, column=0,sticky="W", pady=5, padx=10)

nombre_linea = tk.Label(ventana, text="",font="helvetica 16")
nombre_linea.grid(row=1, column=0, pady=5, padx=10)


nombre_label1= tk.Label(ventana, text="NOMBRE", font="helvetica 12")
nombre_label1.grid(row=2, column=0,sticky="W", pady=5, padx=10)

text_box1 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box1.grid(row=3, column=0,sticky="W", pady=5, padx=10)

nombre_label2= tk.Label(ventana, text="DIRECCION",font="helvetica 12")
nombre_label2.grid(row=2, column=1,sticky="W", pady=5, padx=10)

text_box2 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box2.grid(row=3, column=1,sticky="W", pady=5, padx=10)

nombre_label3= tk.Label(ventana, text="FOLIO",font="helvetica 12")
nombre_label3.grid(row=4, column=0,sticky="W", pady=5, padx=10)

text_box3 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box3.grid(row=5, column=0,sticky="W", pady=5, padx=10)


nombre_label4= tk.Label(ventana, text="FECHA FORM M/D/Y",font="helvetica 12")
nombre_label4.grid(row=4, column=1,sticky="W", pady=5, padx=10)

text_box4 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box4.grid(row=5, column=1,sticky="W", pady=5, padx=10)


nombre_label5= tk.Label(ventana, text="PHONE",font="helvetica 12")
nombre_label5.grid(row=6, column=0,sticky="W", pady=5, padx=10)

text_box5 = tk.Entry(ventana,font="helvetica 12")
text_box5.grid(row=7, column=0,sticky="W", pady=5, padx=10)

nombre_label6= tk.Label(ventana, text="EMAIL",font="helvetica 12")
nombre_label6.grid(row=6, column=1,sticky="W", pady=5, padx=10)

text_box6 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box6.grid(row=7, column=1,sticky="W", pady=5, padx=10)

nombre_label7= tk.Label(ventana, text="PROYECTO",font="helvetica 12")
nombre_label7.grid(row=8, column=0,sticky="W", pady=5, padx=10)

text_box7 = tk.Entry(ventana,font="helvetica 12",validate="key", validatecommand=(validacion, '%P'))
text_box7.grid(row=9, column=0,sticky="W", pady=5, padx=10)

advertencia_label = tk.Label(ventana, fg="red")
advertencia_label.grid(row=50, column=0,sticky="W", pady=5, padx=10)


# Crea una función que imprima el contenido de los cuadros de texto
def imprimir_texto():
    Nombre=text_box1.get()
    Direccion=text_box2.get()
    Folio=text_box3.get()
    Fecha=text_box4.get()
    Phone=text_box5.get()
    Email=text_box6.get()
    Proyecto = text_box7.get()


    if Nombre== "" or Direccion == "" or Folio == "" or Fecha == "" or Phone == "" or Email == "" or Proyecto == "":
        advertencia_label.config(text="Debe llenar todos los campos.",font="helvetica 12")
        messagebox(text="Debe llenar todos los campos.")

    else:
        pass
        
    


    # Copiar el estilo del primer párrafo del documento original al primer párrafo del nuevo documento

    doc.paragraphs[8].add_run(Direccion)
    doc.paragraphs[9].add_run(Folio)
    doc.paragraphs[10].add_run(Nombre)
    doc.paragraphs[14].add_run(Direccion)

    if Email.count("@") :
        doc.paragraphs[16].add_run(Email)


    else:
        advertencia_label.config(text="Debe llenar con correo valido",font="helvetica 12")
        messagebox(text="Debe llenar todos los campos.")
    
    Verificar = len(Phone)
    if Phone.isdigit() and Verificar > 7:

        doc.paragraphs[17].add_run(Phone)

    else:
        advertencia_label.config(text="Debe llenar con Telefono valido",font="helvetica 12")
        messagebox(text="Debe llenar todos los campos.")


    doc.paragraphs[18].add_run(Proyecto)
    doc.paragraphs[19].add_run(Fecha)




    doc.paragraphs[52].add_run(Proyecto)
    


        
    ventana.destroy()

    ventana1 = tk.Tk()
    ventana1.geometry("800x500")
    ventana1.title("Scope of Work")
    ventana1.resizable(False, False)

    nombre_label8= tk.Label(ventana1, text="Scope of Work",font="helvetica 16")
    nombre_label8.grid(row=0, column=0,sticky="W")

    nombre_Concepto= tk.Label(ventana1, text="Concepto:",font="helvetica 11")
    nombre_Concepto.grid(row=0, column=2,sticky="W")




    def checkbox_changed(checkbox1, checkbox2):
        if checkbox1.get():
            checkbox2.set(False)
        elif checkbox2.get():
            checkbox1.set(False)



    #NOC
    var1 =tk.IntVar()
    c = tk.Checkbutton( ventana1, text="SI",font="helvetica 11", variable=var1, command=lambda: checkbox_changed(var1,var15))
    c.grid(row=1, column=0,sticky="W")
    #NOC
    var15 = tk.IntVar()
    c15 = tk.Checkbutton(ventana1, text="NO", font="helvetica 11",variable=var15, command=lambda: checkbox_changed(var15, var1))
    c15.grid(row=1, column=1,sticky="W")

    LABEL_NOC = tk.Label(ventana1, text=" NOC",font="helvetica 11" )
    LABEL_NOC.grid(row=1, column=2,sticky="W", pady=5,padx=10)



    #LABOR
    var2 =tk.IntVar()
    c2= tk.Checkbutton( ventana1, text="SI",font="helvetica 11", variable=var2,command=lambda: checkbox_changed(var2, var16))
    c2.grid(row=2, column=0,sticky="W")

    var16 = tk.IntVar()
    c16 = tk.Checkbutton(ventana1, text="NO", font="helvetica 11",variable=var16, command=lambda: checkbox_changed(var16, var2))
    c16.grid(row=2, column=1,sticky="W")

    LABEL_LABOR = tk.Label(ventana1, text="LABOR",font="helvetica 11" )
    LABEL_LABOR.grid(row=2, column=2,sticky="W", pady=5,padx=10)



    #PERMIT FEE
    var3 =tk.IntVar()
    c3= tk.Checkbutton( ventana1, text="SI",font="helvetica 11", variable=var3, command=lambda: checkbox_changed(var3, var12))
    c3.grid(row=8, column=0,sticky="W")

    var12 = tk.IntVar()
    c12 = tk.Checkbutton(ventana1, text="NO", font="helvetica 11",variable=var12, command=lambda: checkbox_changed(var12, var3) )
    c12.grid(row=8, column=1,sticky="W")

    LABEL_FEE = tk.Label(ventana1, text="PERMIT FEE",font="helvetica 11" )
    LABEL_FEE.grid(row=8, column=2,sticky="W", pady=5,padx=10)





    #Calificacion buiLding
    var4 =tk.IntVar()
    c4= tk.Checkbutton( ventana1, text="SI",font="helvetica 11", variable=var4, command=lambda: checkbox_changed(var4, var11))
    c4.grid(row=4, column=0,sticky="W")

    var11= tk.IntVar()
    c11 = tk.Checkbutton(ventana1, text="NO", font="helvetica 11", variable=var11, command=lambda: checkbox_changed(var11, var4))
    c11.grid(row=4, column=1,sticky="W")

    LABEL_BU = tk.Label(ventana1, text="CALIFICACION DE: ",font="helvetica 11" )
    LABEL_BU.grid(row=4, column=2,sticky="W", pady=5,padx=10)

    Text_CALI = tk.Entry(ventana1, font="helvetica 12",)
    Text_CALI.grid(row=4, column=3,sticky="w", pady=5,padx=5)




        
    #Tramite con la ciudad
    var8=tk.IntVar()
    c8= tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var8, command=lambda: checkbox_changed(var8, var20))
    c8.grid(row=5, column=0,sticky="w")

    var20 =tk.IntVar()
    b1 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var20, command=lambda: checkbox_changed(var20, var8))
    b1.grid( row=5, column=1,sticky="w" )

    LABEL_CID = tk.Label(ventana1, text="TRAMITES CON LA CIUDAD",font="helvetica 11" )
    LABEL_CID.grid(row=5, column=2,sticky="w", pady=5,padx=10)


    #Survey
    var9=tk.IntVar()
    c9= tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var9, command=lambda: checkbox_changed(var9, var21))
    c9.grid(row=6, column=0,sticky="W")

    var21 =tk.IntVar()
    b2 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var21, command=lambda: checkbox_changed(var21, var9))
    b2.grid( row=6, column=1,sticky="w" )

    LABEL_SV = tk.Label(ventana1, text="SURVEY",font="helvetica 11" )
    LABEL_SV.grid(row=6, column=2,sticky="w", pady=5,padx=10)


    #Material
    var22 =tk.IntVar()
    b3 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var22, command=lambda: checkbox_changed(var22, var23))
    b3.grid( row=3, column=0,sticky="w" )

    var23 =tk.IntVar()
    b4 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var23, command=lambda: checkbox_changed(var23, var22))
    b4.grid( row=3, column=1,sticky="w" )

    LABEL_MT = tk.Label(ventana1, text="MATERIAL",font="helvetica 11" )
    LABEL_MT.grid(row=3, column=2,sticky="w", pady=5,padx=10)


    #Planos
    var24 =tk.IntVar()
    b5 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var24, command=lambda: checkbox_changed(var24, var25))
    b5.grid( row=7, column=0,sticky="w" )

    var25 =tk.IntVar()
    b6 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var25, command=lambda: checkbox_changed(var25, var24))
    b6.grid( row=7, column=1,sticky="w" )

    LABEL_MT = tk.Label(ventana1, text="PLANOS DE:",font="helvetica 11" )
    LABEL_MT.grid(row=7, column=2,sticky="w", pady=5,padx=10)

    Text_Planos = tk.Entry(ventana1, font="helvetica 12")
    Text_Planos.grid(row=7, column=3,sticky="w", pady=5,padx=5)



    var26 =tk.IntVar()
    b7 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var26, command=lambda: checkbox_changed(var26, var27))
    b7.grid( row=9, column=0,sticky="w" )

    var27 =tk.IntVar()
    b8 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var27, command=lambda: checkbox_changed(var27, var26))
    b8.grid( row=9, column=1,sticky="w" )

    LABEL_MNT = tk.Label(ventana1, text="NUEVO CONCEPTO",font="helvetica 11" )
    LABEL_MNT.grid(row=9, column=2,sticky="w", pady=5,padx=10)



    var28 =tk.IntVar()
    b9 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var28, command=lambda: checkbox_changed(var28, var29))
    b9.grid( row=10, column=0,sticky="w" )

    var29 =tk.IntVar()
    b10 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var29, command=lambda: checkbox_changed(var29, var28))
    b10.grid( row=10, column=1,sticky="w" )

    LABEL_MMT = tk.Label(ventana1, text="NUEVO CONCEPTO",font="helvetica 11" )
    LABEL_MMT.grid(row=10, column=2,sticky="w", pady=5,padx=10)




    var30 =tk.IntVar()
    b11 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var30, command=lambda: checkbox_changed(var30, var31))
    b11.grid( row=11, column=0,sticky="w" )

    var31 =tk.IntVar()
    b12 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var31 , command=lambda: checkbox_changed(var31, var30))
    b12.grid( row=11, column=1,sticky="w" )

    LABEL_NT = tk.Label(ventana1, text="NUEVO CONCEPTO",font="helvetica 11" )
    LABEL_NT.grid(row=11, column=2,sticky="w", pady=5,padx=10)
    



    var32 =tk.IntVar()
    b13 = tk.Checkbutton(ventana1, text="SI",font="helvetica 11", variable=var32, command=lambda: checkbox_changed(var32, var33))
    b13.grid( row=12, column=0,sticky="w" )

    var33 =tk.IntVar()
    b14 = tk.Checkbutton(ventana1, text="NO",font="helvetica 11", variable=var33, command=lambda: checkbox_changed(var33, var32))
    b14.grid( row=12, column=1,sticky="w" )

    LABEL_TT = tk.Label(ventana1, text="NUEVO CONCEPTO",font="helvetica 11" )
    LABEL_TT.grid(row=12, column=2,sticky="w", pady=5,padx=10)

    Text_Concepto = tk.Entry(ventana1, font="helvetica 12")
    Text_Concepto.grid(row=12, column=3, sticky="w", pady=5, padx=5)




    def imprimir():



        # buscaremos capturar el valor del checkbox para poder compararlo y decidir que introducir al contrato.
        a = var15.get()
        if a == 1:
            doc.paragraphs[53].add_run("Owner is responsible for recording the Notice of Commencement with the court")

        aa = var1.get()
        if aa == 1:
            doc.paragraphs[53].add_run("This proposal includes the filing of a Notice of Commencement with the Court (NOC)")
        
   
        b = var16.get()
        if b == 1:
            doc.paragraphs[54].add_run("Proposal does not include any form of labor or payment of any subcontractors, any materials or any expenses not explicitly included within the scope of work. ")

        bb = var2.get()
        if bb == 1:
               doc.paragraphs[54].add_run("xxx labor si ") 

    
        c = var12.get()
        cc = var3.get()

        if c == 1:
            doc.paragraphs[55].add_run("Owner is responsible for the cost of all permits. This proposal does not include the cost of permits, municipal fees, application fees, or any other city fees. Owner is responsible for any and all fees. ")

        if cc == 1:
            doc.paragraphs[55].add_run("xxxx")

 
        
        d = var4.get()#buil
        # Los no
        z = var11.get()#buil
        DD = Text_CALI.get()
        if d == 1 and DD =="":
             LABEL_BU.config(text="FALTA CALIFICACION.", fg="red", font="helvetica 12")
             messagebox(text="Debe llenar todos los campos.")

        if z==1:
            doc.paragraphs[56].add_run("This proposal does not include the use of any trade license (Building, electrical, plumbing, roofing, mechanical, etc.). Should the Owner require any trades, Contractor will provide a separate proposal for an additional fee. ")
            
        if d == 1:
            doc.paragraphs[56].add_run(f"This proposal includes preparation of the {DD} permit applications. This proposal does not include any other commercial licenses or work that isn’t detailed above. ")


        h = var24.get()
        hhh = var25.get()
        hh = Text_Planos.get()
        
        if h == 1 and hh =="":
            LABEL_MT.config(text="FALTA ESPECIFICAR.", fg="red", font="helvetica 12")
            messagebox("FALTA ESPECIFICAR")


        


        if h == 1:
            doc.paragraphs[57].add_run(f"Proposal includes price for drawings that will be submitted to the city, drawings will be prepared and provided to Owner. These drawings will depict only the {hh} .The rest of the property will not be depicted in the drawings.  ")
        
        if hhh == 1:
            doc.paragraphs[57].add_run("xxxx")

            

        i = var21.get()
        if i == 1:
            doc.paragraphs[58].add_run("A survey is not included as a part of this contract.")
        ii = var9.get()
        if ii == 1:
             doc.paragraphs[58].add_run("A survey is included as a part of this contract.")


        

        
        n = var26.get()
        nn = var27.get()
        if n ==1:
            doc.paragraphs[59].add_run("")
        if nn ==1:
            doc.paragraphs[59].add_run("")


        m = var28.get()
        mm = var29.get()
        if m ==1:
            doc.paragraphs[60].add_run("")
        if mm ==1:
            doc.paragraphs[60].add_run("")
        
        w = var30.get()
        ww = var31.get()
        if w ==1:
            doc.paragraphs[61].add_run("")
        if ww ==1:
            doc.paragraphs[61].add_run("")
        

        v = var22.get()
        vv = var23.get()
        if v == 1:
             doc.paragraphs[62].add_run("")
        if vv == 1:
             doc.paragraphs[62].add_run("")


        Imprimir = Text_Concepto.get()
        f = var32.get()
        ff = var33.get()
        if f ==1:
            doc.paragraphs[63].add_run(Imprimir)
        if ff ==1:
            doc.paragraphs[63].add_run("")


            

            



        ventana1.destroy()

        ventana2 = tk.Tk()
        ventana2.geometry("500x400")
        ventana2.title("Agenda de Pagos.")

        contenedor = tk.Frame(ventana2)
        contenedor.grid(row=0, column=1,sticky="W", pady=5, padx=10)

        contenedor3 = tk.Frame(ventana2)
        contenedor3.grid(row=10, column=0,sticky="W", pady=5, padx=10)

        contenedor2 = tk.Frame(ventana2)
        contenedor2.grid(row=0, column=0,sticky="W", pady=5, padx=10)

        nombre_label18= tk.Label(contenedor2, text="Costo del proyecto",font="helvetica 16")
        nombre_label18.grid(row=0, column=0,sticky="W", pady=5, padx=10)

        nombre_label21 = tk.Label(contenedor2,text="Monto del contrato",font="helvetica 12" )
        nombre_label21.grid(row=1, column=0,sticky="W", pady=5, padx=10)
        text_box21 = tk.Entry(contenedor2,font="helvetica 12")
        text_box21.grid(row=2, column=0,sticky="W", pady=5, padx=10)


        nombre_label22 = tk.Label(contenedor2,text="PORCIENTO DEL PRIMER PAGO",font="helvetica 12" )
        nombre_label22.grid(row=3, column=0,sticky="W", pady=5, padx=10)

        text_box22 = tk.Entry(contenedor2,font="helvetica 12")
        text_box22.grid(row=4, column=0,sticky="W", pady=5, padx=10)



        nombre_label23 = tk.Label(contenedor2,text="PORCIENTO DEL SEGUNDO PAGO",font="helvetica 12" )
        nombre_label23.grid(row=5, column=0,sticky="W", pady=5, padx=10)

        text_box23 = tk.Entry(contenedor2,font="helvetica 12")
        text_box23.grid(row=6, column=0,sticky="W", pady=5, padx=10)



        nombre_label24 = tk.Label(contenedor2,text="PORCIENTO DEL TERCER PAGO",font="helvetica 12" )
        nombre_label24.grid(row=7, column=0,sticky="W", pady=5, padx=10)

        text_box24 = tk.Entry(contenedor2,font="helvetica 12")
        text_box24.grid(row=8, column=0,sticky="W", pady=5, padx=10)


        advertencia_label = tk.Label(contenedor3, fg="red")
        advertencia_label.grid(row=9, column=0,sticky="W", pady=5, padx=10)






        def finalizar():
            por1 = text_box22.get()
            por2 = text_box23.get()
            por3 = text_box24.get()


            if por1 == "" or por2 == "" or por3 == "":
                advertencia_label.config(text="No puede usar esta opcion sin llenar todos los campos.",font="helvetica 12")
            else:
                advertencia_label.config(text="")



            Porciento1 = float("0." + por1)
            Porciento2 = float("0." + por2)
            Porciento3 = float("0." + por3)

            v = text_box21.get()

            if v == "":

                advertencia_label.config(text="No puedes usar esta opcion sin llenar todos los campos.",font="helvetica 12")


            Pago = int(v)

            doc.paragraphs[68].add_run("Payment schedule:")
            

            doc.paragraphs[69].add_run("Payment for the above quoted work will be due as follows:")
            

            doc.paragraphs[70].add_run(f"Total Proposal/Contract price: ${Pago}.00 (Based on scope of work shown above)")
            


            
            doc.paragraphs[71].add_run(f"The amount of ${Pago * Porciento1} of the contract price will be due upon signing the contract.")
            
            
            doc.paragraphs[72].add_run(f"The amount of ${ Pago * Porciento2} will be due after rough inspection")
           
            
            doc.paragraphs[73].add_run(f"The amount of ${Pago * Porciento3} will be due prior to requesting final inspections")
            







            doc.save("Nuevo_Contrato " + Nombre +".Docx")
            ventana2.destroy()

        def One():
            texto = text_box21.get()
            if not texto:
                advertencia_label.config(text="El monto del contrato no puede estar vacio",font="helvetica 12" )
            else:
                advertencia_label.config(text="")


            Pago = int(text_box21.get())

                

            doc.paragraphs[68].add_run("Payment schedule:")
            

            doc.paragraphs[69].add_run("Payment for the above quoted work will be due as follows:")
            

            doc.paragraphs[70].add_run(f"The deposit will total the amount of ${Pago} and is due upon Owner's Signature ")
            

            doc.save("Nuevo_Contrato " + Nombre +".Docx")
            ventana2.destroy()


        def Cuota():
            v = text_box21.get()
            if v == "":
                messagebox.showerror("Error", "El Monto del contrato no puede estar Vacio.")

            Pago = int(v)

            doc.paragraphs[68].add_run("Payment schedule:")
            

            doc.paragraphs[69].add_run("Payment for the above quoted work will be due as follows:")
            

            doc.paragraphs[70].add_run(f"Total Proposal/Contract price: ${Pago}.00 (Based on scope of work shown above)")
            


            
            doc.paragraphs[71].add_run(f"The upfront deposit will total the amount of ${Pago * 0.5} is due upon Owner's Signature. No work will be provided or started until contractor has received initial payment.")
            
            
            doc.paragraphs[72].add_run(f"The second payment of ${Pago * 0.5} is due upon approval of building permit. No work will start/continue until contractor has received the second payment.")
            







            doc.save("Nuevo_Contrato " + Nombre +".Docx")
            ventana2.destroy()


        
        
        
        
        
        
        
        
        boton1 = tk.Button(contenedor, text="3 Pagos por %",command=finalizar, width=15,height=2, font="helvetica 12")
        boton1.grid(row=0, column=0,sticky="W", pady=5, padx=10)


        boton2 = tk.Button(contenedor, text="Un Pago",command=One, width=15,height=2, font="helvetica 12")
        boton2.grid(row=1, column=0,sticky="W", pady=5, padx=10)


        boton2 = tk.Button(contenedor, text="2 Cuotas",command=Cuota, width=15,height=2, font="helvetica 12")
        boton2.grid(row=2, column=0,sticky="W", pady=5, padx=10)





        


    
    boton = tk.Button(ventana1, text="Siguiente", font="helvetica 12", command=imprimir, width=15, height=2 )
    boton.grid(row=17, column=3,sticky="W")











    ventana1.mainloop()
    

    # Guardar el nuevo documento
def on_enter(event):
    boton.config(bg='#D0BDBD', width=16,height=3)

def on_leave(event):
    boton.config(bg='white',width=15,height=2)   



# Crea un botón para llamar a la función
boton = tk.Button(ventana, text="Siguiente", command=imprimir_texto, width=15,height=2, font="helvetica 12",bg='white')
boton.grid(row=50, column=1,sticky="W")

boton.bind("<Enter>", on_enter)
boton.bind("<Leave>", on_leave)

# Inicia el bucle de eventos de tkinter
ventana.mainloop()